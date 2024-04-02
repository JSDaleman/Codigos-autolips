import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Leer datos del archivo de Excel
def leer_excel(nombre_archivo):
    wb = load_workbook(nombre_archivo)
    hoja = wb.active
    destinatarios = []
    for fila in hoja.iter_rows(min_row=3, values_only=True):
        if all(celda is None for celda in fila):  # Verifica si todas las celdas de la fila están vacías
           break  # Detiene la iteración si la fila está vacía
        destinatarios.append({
            'nombre': fila[1],
            'correo': fila[7]
        })
    return destinatarios

# Ruta donde deseas guardar los documentos de Word
ruta_guardado = r"D:\Escritorio\Proyectos U\Semestre10\ComputacionGrafica\Trabajos\AutomatizacionV2\Documentos"

# Obtener la ruta absoluta del archivo
current_dir = os.path.dirname(__file__)
excel_file = os.path.join(current_dir, 'Datos.xlsx')

# Leer datos del archivo Excel
destinatarios = leer_excel(excel_file)

# Escribir correo en Word
def escribir_correo(destinatario):
    documento = Document()
    documento.add_heading('Correo', level=1)
    p = documento.add_paragraph()
    p.add_run('Estimado/a ' + destinatario['nombre'] + ',').bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    documento.add_paragraph('Este es un correo de ejemplo.')
    
    documento.add_paragraph().add_run('Atentamente,').italic = True
    documento.add_paragraph('Area administrativa')
    
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Guardar el documento en la ruta especificada
    nombre_archivo = f"correo_para_{destinatario['nombre']}.docx"
    ruta_completa = os.path.join(ruta_guardado, nombre_archivo)
    documento.save(ruta_completa)
# Crear y guardar correos para cada destinatario
for destinatario in destinatarios:
    escribir_correo(destinatario)

print("Los correos han sido creados exitosamente.")
